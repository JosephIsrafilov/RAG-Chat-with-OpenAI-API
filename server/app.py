import os
import io
import tempfile
from typing import List, Tuple
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import OpenAI
import numpy as np
import faiss
import tiktoken
from pypdf import PdfReader
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

EMBED_MODEL = "text-embedding-3-large"
GEN_MODEL = "gpt-4.1"
CHUNK_TOKENS = 400
CHUNK_OVERLAP = 60
TOP_K_DEFAULT = 6

try:
    import textract 
    HAS_TEXTRACT = True
except Exception:
    HAS_TEXTRACT = False

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

docs: List[Tuple[str, str]] = []
embs: np.ndarray | None = None
index: faiss.Index | None = None

def read_file_to_text(file_bytes: bytes, name: str) -> str:
    n = name.lower()
    if n.endswith(".txt") or n.endswith(".md"):
        return file_bytes.decode("utf-8", errors="ignore")
    if n.endswith(".pdf"):
        pdf = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for p in pdf.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n\n".join(pages)
    if n.endswith(".docx"):
        from docx import Document
        d = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in d.paragraphs if p.text]
        for table in d.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells if cell.text))
        return "\n".join(parts)
    if n.endswith(".doc"):
        if not HAS_TEXTRACT:
            return ""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            return textract.process(tmp_path).decode("utf-8", errors="ignore")
        finally:
            os.unlink(tmp_path)
    return ""

def chunk_text(text: str) -> List[str]:
    enc = tiktoken.get_encoding("cl100k_base")
    tokens = enc.encode(text)
    chunks = []
    i = 0
    while i < len(tokens):
        chunk = tokens[i:i+CHUNK_TOKENS]
        chunks.append(enc.decode(chunk))
        i += max(1, CHUNK_TOKENS - CHUNK_OVERLAP)
    return [c.strip() for c in chunks if c.strip()]

def embed_texts(texts: List[str]) -> np.ndarray:
    B = 1000
    vecs = []
    for i in range(0, len(texts), B):
        batch = texts[i:i+B]
        resp = client.embeddings.create(model=EMBED_MODEL, input=batch)
        vecs.extend([e.embedding for e in resp.data])
    return np.array(vecs, dtype="float32")

def build_faiss(embeddings: np.ndarray) -> faiss.Index:
    faiss.normalize_L2(embeddings)
    idx = faiss.IndexFlatIP(embeddings.shape[1])
    idx.add(embeddings)
    return idx

def search(q_vec: np.ndarray, top_k: int):
    v = q_vec.copy()
    faiss.normalize_L2(v)
    D, I = index.search(v, top_k)
    return D, I

def make_prompt(question: str, hits: List[Tuple[str, str]]) -> list[dict]:
    sources_text = "\n\n".join([f"[{i+1}] ({src})\n{txt}" for i,(src,txt) in enumerate(hits)])
    system = (
        "You are a helpful RAG assistant. "
        "Answer the user's question using ONLY the provided context if possible. "
        "If the answer is not in the context, say you don't have enough information. "
        "Cite sources as [#] where # is the context index."
    )
    user = (
        f"Question:\n{question}\n\n"
        f"Context chunks:\n{sources_text}\n\n"
        "Instructions:\n- If you use multiple chunks, cite like [1][3].\n- Be concise and precise."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]

class AskPayload(BaseModel):
    question: str
    top_k: int | None = None

@app.post("/upload")
async def upload(files: List[UploadFile] = File(...)):
    global docs
    added = 0
    for f in files:
        b = await f.read()
        text = read_file_to_text(b, f.filename)
        if not text.strip():
            continue
        for c in chunk_text(text):
            docs.append((f.filename, c))
            added += 1
    return {"status": "ok", "files": len(files), "chunks_added": added, "total_chunks": len(docs)}

@app.post("/build")
async def build():
    global embs, index
    if not docs:
        return {
            "status": "ok",
            "chunks": 0,
            "message": "No documents to index. Please upload files and try again."
        }
    texts = [c for _, c in docs]
    embs = embed_texts(texts)
    index = build_faiss(embs)
    return {"status": "ok", "chunks": len(texts)}


@app.post("/ask")
async def ask(payload: AskPayload):
    if index is None or not docs:
        return {
            "status": "ok",
            "answer": (
                "I don't have enough information to answer. "
                "Please upload documents and click Build Index."
            ),
            "sources": []
        }

    if not payload.question.strip():
        return {"status": "no_question"}

    q_emb = client.embeddings.create(
        model=EMBED_MODEL,
        input=[payload.question]
    ).data[0].embedding
    q_emb = np.array([q_emb], dtype="float32")

    top_k = payload.top_k or TOP_K_DEFAULT
    top_k = max(1, min(top_k, len(docs))) 

    D, I = search(q_emb, top_k=top_k)
    idxs = I[0].tolist()
    hits = [(docs[j][0], docs[j][1]) for j in idxs]

    messages = make_prompt(payload.question, hits)
    completion = client.chat.completions.create(
        model=GEN_MODEL,
        messages=messages,
        temperature=0.2
    )
    answer = completion.choices[0].message.content

    sources = [
        {"id": i + 1, "file": s, "preview": t[:300]}
        for i, (s, t) in enumerate(hits)
    ]
    return {"status": "ok", "answer": answer, "sources": sources}


@app.post("/reset")
async def reset():
    global docs, embs, index
    docs = []
    embs = None
    index = None
    return {"status": "ok"}
